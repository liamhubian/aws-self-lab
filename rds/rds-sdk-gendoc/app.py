import boto3
from docx import Document
from docx.shared import Inches


rds_client = boto3.client('rds')
document = Document()

def main():
    db_instances = rds_client.describe_db_instances()["DBInstances"]

    for db_instance in db_instances:
        instance_info = "======" \
                + "\ninstance identifier    : " + db_instance["DBInstanceIdentifier"] \
                + "\ndatabase engine        : " + db_instance["Engine"] \
                + "\ndatabase endpoint      : " + db_instance["Endpoint"]["Address"] + ":" + str(db_instance["Endpoint"]["Port"])
        print(instance_info) 

    document.add_heading('RDS Information', level=0)

    rds_table = document.add_table(rows=1, cols=3)
    rds_table_widths = (Inches(0.4), Inches(1), Inches(3))

    head_cells = rds_table.rows[0].cells
    head_cells[0].text = "Number"
    head_cells[1].text = "Instance"
    head_cells[2].text = "Description"

    rds_instance_count = 0
    for db_instance in db_instances:
        rds_instance_count+=1
        row_cells = rds_table.add_row().cells

        instance_description = \
                "engine: " + db_instance["Engine"] \
                + "\ndatabase endpoint: " + db_instance["Endpoint"]["Address"] + ":" + str(db_instance["Endpoint"]["Port"])

        row_cells[0].text = str(rds_instance_count)
        row_cells[1].text = db_instance["DBInstanceIdentifier"]
        row_cells[2].text = instance_description

    set_table_width(rds_table, rds_table_widths)

    document.add_page_break()
    document.save("rds.docx")

def set_table_width(table, widths):
    for row in table.rows:
        for cell_index, width in enumerate(widths):
            row.cells[cell_index].width = width

                                
if __name__ == "__main__":      
    main()
